"""Article export: turning an Article's raw content into a downloadable file for a
chosen format (ADR-0007). Single choke point (`build_export_filename`/`build_export_document`)
shared by the initial delivery-on-generation path and the history re-download path (#20).
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.text.paragraph import Paragraph

from .types import ArticleFormat, ArticleView


def build_export_filename(title: str, platform: str, article_format: ArticleFormat) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-") or "article"
    return f"{slug}-{platform}.{article_format}"


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MAX_HEADING_LEVEL = (
    4  # python-docx's default template only styles Heading 1-4 distinctly from Normal
)


def build_export_document(article: ArticleView, article_format: ArticleFormat) -> bytes:
    text = article.content.decode("utf-8")
    if article_format == "md":
        return text.encode("utf-8")
    document = Document()
    for line in text.split("\n"):
        heading_match = _HEADING_RE.match(line)
        bullet_match = _BULLET_RE.match(line)
        if heading_match:
            level = min(len(heading_match.group(1)), _MAX_HEADING_LEVEL)
            _add_paragraph_with_bold_runs(
                document.add_heading("", level=level), heading_match.group(2)
            )
        elif bullet_match:
            _add_paragraph_with_bold_runs(
                document.add_paragraph(style="List Bullet"), bullet_match.group(1)
            )
        else:
            _add_paragraph_with_bold_runs(document.add_paragraph(), line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_paragraph_with_bold_runs(paragraph: Paragraph, text: str) -> None:
    position = 0
    for match in _BOLD_RE.finditer(text):
        paragraph.add_run(text[position : match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    paragraph.add_run(text[position:])
