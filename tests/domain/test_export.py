from io import BytesIO

from docx import Document

from content_zavod.domain import (
    ArticleId,
    ArticleView,
    PlanItemId,
    build_export_document,
    build_export_filename,
)

_ARTICLE = ArticleView(
    id=ArticleId("article-1"),
    plan_item_id=PlanItemId("item-1"),
    title="Как выбрать нишу!",
    platform="zen",
    content="Первый абзац.\n\nВторой абзац.".encode(),
)


def test_build_export_filename_slugifies_title_and_appends_format() -> None:
    assert build_export_filename("Best Niche Guide", "zen", "docx") == "best-niche-guide-zen.docx"
    assert build_export_filename("Best Niche Guide", "zen", "md") == "best-niche-guide-zen.md"


def test_build_export_filename_falls_back_to_article_for_an_empty_slug() -> None:
    assert build_export_filename("!!!", "zen", "docx") == "article-zen.docx"


def test_build_export_document_md_is_the_plain_text_content() -> None:
    assert build_export_document(_ARTICLE, "md") == _ARTICLE.content


def test_build_export_document_docx_contains_the_content_as_paragraphs() -> None:
    document_bytes = build_export_document(_ARTICLE, "docx")

    document = Document(BytesIO(document_bytes))
    paragraphs = [p.text for p in document.paragraphs]

    assert paragraphs == ["Первый абзац.", "", "Второй абзац."]


def test_build_export_document_docx_renders_markdown_headings_as_word_headings() -> None:
    article = ArticleView(
        id=_ARTICLE.id,
        plan_item_id=_ARTICLE.plan_item_id,
        title=_ARTICLE.title,
        platform=_ARTICLE.platform,
        content="# Заголовок\n\nАбзац.".encode(),
    )

    document = Document(BytesIO(build_export_document(article, "docx")))

    heading = document.paragraphs[0]
    assert heading.text == "Заголовок"
    assert heading.style.name == "Heading 1"


def test_build_export_document_docx_renders_bold_inside_a_heading() -> None:
    article = ArticleView(
        id=_ARTICLE.id,
        plan_item_id=_ARTICLE.plan_item_id,
        title=_ARTICLE.title,
        platform=_ARTICLE.platform,
        content="## Что **важно** знать".encode(),
    )

    document = Document(BytesIO(build_export_document(article, "docx")))
    heading = document.paragraphs[0]

    assert heading.text == "Что важно знать"
    assert heading.style.name == "Heading 2"
    bold_runs = [r for r in heading.runs if r.bold]
    assert [r.text for r in bold_runs] == ["важно"]


def test_build_export_document_docx_renders_markdown_bullets_as_list_items() -> None:
    article = ArticleView(
        id=_ARTICLE.id,
        plan_item_id=_ARTICLE.plan_item_id,
        title=_ARTICLE.title,
        platform=_ARTICLE.platform,
        content="- Первый пункт\n- Второй пункт".encode(),
    )

    document = Document(BytesIO(build_export_document(article, "docx")))
    paragraphs = document.paragraphs

    assert [p.text for p in paragraphs] == ["Первый пункт", "Второй пункт"]
    assert all(p.style.name == "List Bullet" for p in paragraphs)


def test_build_export_document_docx_renders_markdown_bold_as_a_bold_run() -> None:
    article = ArticleView(
        id=_ARTICLE.id,
        plan_item_id=_ARTICLE.plan_item_id,
        title=_ARTICLE.title,
        platform=_ARTICLE.platform,
        content="Обычный текст с **важным** словом.".encode(),
    )

    document = Document(BytesIO(build_export_document(article, "docx")))
    runs = document.paragraphs[0].runs

    bold_runs = [r for r in runs if r.bold]
    assert [r.text for r in bold_runs] == ["важным"]
    assert document.paragraphs[0].text == "Обычный текст с важным словом."
